# app/core/resume_parser.py - Enhanced Resume Parser
import PyPDF2
import docx
import re
from typing import Dict, List
import nltk
from nltk.corpus import stopwords
import os
import tempfile

# Try to import optional libraries
try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# Ensure NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)

class ResumeParser:
    def __init__(self):
        self.skills_db = self._load_skills_db()
    
    def _load_skills_db(self):
        """Comprehensive skills database"""
        return {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust', 'php', 'swift', 'kotlin', 'typescript', 'html', 'css'],
            'frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'node.js', 'express', 'spring', 'laravel', 'rails'],
            'databases': ['sql', 'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'firebase', 'dynamodb'],
            'cloud_devops': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'ci/cd', 'terraform'],
            'ai_ml': ['tensorflow', 'pytorch', 'scikit-learn', 'keras', 'pandas', 'numpy', 'nltk', 'spacy', 'opencv'],
            'design': ['figma', 'canva', 'adobe xd', 'photoshop', 'illustrator', 'ui/ux', 'web design', 'responsive design'],
            'soft_skills': ['leadership', 'communication', 'teamwork', 'problem solving', 'project management', 'agile', 'scrum']
        }
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            return text
        except Exception as e:
            return ""
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF"""
        text = ""
        if HAS_FITZ:
            try:
                pdf_file.seek(0)
                doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
                return text
            except:
                pass
        return text
    
    def extract_text_from_txt(self, txt_file) -> str:
        """Extract text from TXT"""
        try:
            return txt_file.read().decode('utf-8')
        except:
            return ""
    
    # ============ ENHANCED EXTRACTION METHODS ============
    
    def extract_email(self, text: str) -> str:
        """Extract email address"""
        patterns = [
            r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
            r'Email:\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0) if '@' in match.group(0) else match.group(1)
        return None
    
    def extract_phone(self, text: str) -> str:
        """Extract phone number"""
        patterns = [
            r'\+?[\d\s\-\(\)]{10,20}',
            r'Phone:\s*([\+]?[\d\s\-\(\)]{10,20})',
            r'Tel:\s*([\+]?[\d\s\-\(\)]{10,20})'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                digits = re.sub(r'\D', '', match if isinstance(match, str) else str(match))
                if len(digits) >= 10:
                    return match.strip() if isinstance(match, str) else str(match).strip()
        return None
    
    def extract_linkedin(self, text: str) -> str:
        """Extract LinkedIn URL"""
        patterns = [
            r'linkedin\.com/in/[\w\-]+',
            r'linkedin\.com/company/[\w\-]+'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0)
        return None
    
    def extract_location(self, text: str) -> str:
        """Extract location"""
        patterns = [
            r'Location:\s*([A-Za-z\s,]+)',
            r'Based in:\s*([A-Za-z\s,]+)',
            r'Address:\s*([A-Za-z\s,]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def extract_summary(self, text: str) -> str:
        """Extract professional summary"""
        patterns = [
            r'Professional Summary[:\s]*([^\n]+(?:\n[^\n]+)*?)(?=\n\n|\n[A-Z])',
            r'Summary[:\s]*([^\n]+(?:\n[^\n]+)*?)(?=\n\n|\n[A-Z])',
            r'Profile[:\s]*([^\n]+(?:\n[^\n]+)*?)(?=\n\n|\n[A-Z])'
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                summary = match.group(1).strip().replace('\n', ' ')[:500]
                return summary
        return None
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract all skills from text"""
        skills = set()
        text_lower = text.lower()
        
        for category, category_skills in self.skills_db.items():
            for skill in category_skills:
                if skill in text_lower:
                    skills.add(skill.title())
        
        # Also look for skills section
        skills_section = re.search(r'Skills?[:\s]*([^\n]+(?:\n[^\n]+)*?)(?=\n\n|\n[A-Z])', text, re.IGNORECASE)
        if skills_section:
            skill_text = skills_section.group(1)
            # Extract bullet points or comma-separated skills
            skill_items = re.findall(r'[•\-]\s*([A-Za-z\s/&]+)', skill_text)
            if not skill_items:
                skill_items = re.split(r'[,\n]', skill_text)
            for item in skill_items:
                cleaned = item.strip()
                if len(cleaned) > 2 and len(cleaned) < 50:
                    skills.add(cleaned)
        
        return list(skills)[:20]
    
    def extract_experience_years(self, text: str) -> int:
        """Extract years of experience"""
        # Look for date ranges
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4}|Present)',
            r'(\d{4})\s*to\s*(\d{4}|present)',
            r'From\s*(\d{4})\s*to\s*(\d{4}|present)'
        ]
        
        years = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                start = int(match[0])
                end = 2026 if match[1].lower() in ['present', 'current'] else int(match[1])
                if start and end:
                    years.append(end - start)
        
        if years:
            return max(min(sum(years), 15), 1)
        
        # Look for explicit statement
        exp_patterns = [
            r'(\d+)\+?\s*years?\s+of\s+experience',
            r'experience\s+of\s+(\d+)\+?\s*years?'
        ]
        for pattern in exp_patterns:
            match = re.search(pattern, text.lower())
            if match:
                return int(match.group(1))
        
        return 2
    
    def extract_education(self, text: str) -> List[Dict]:
        """Extract education details"""
        education = []
        patterns = [
            r'(B\.?S\.?|B\.?A\.?|Bachelor)[^\n]*?(?:in\s+)?([^\n]+)',
            r'(M\.?S\.?|M\.?A\.?|Master)[^\n]*?(?:in\s+)?([^\n]+)',
            r'(Ph\.?D|Doctorate)[^\n]*?(?:in\s+)?([^\n]+)',
            r'(BSc|MSc|BA|MA|BBA|MBA)[^\n]*?(?:in\s+)?([^\n]+)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                degree = match[0].strip() if isinstance(match, tuple) else match.strip()
                field = match[1].strip() if len(match) > 1 else ""
                if len(degree) > 2 and degree not in [e.get('degree') for e in education]:
                    education.append({
                        "degree": degree,
                        "field": field,
                        "institution": self._extract_institution(text, degree)
                    })
        
        return education[:3]
    
    def _extract_institution(self, text: str, degree: str) -> str:
        """Extract institution name near degree"""
        pattern = f"{degree}[^\n]*?(?:from|at|,)\\s*([A-Za-z\\s]+(?:University|College|Institute|School))"
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
        return "Not specified"
    
    def extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certs = []
        cert_patterns = [
            r'(?:Certified|Certificate)[:\s]*([A-Za-z\s&]+)',
            r'(AWS|Azure|Google|Microsoft|Cisco)\s+(?:Certified|Certification)',
            r'(PMP|CISSP|CEH|CISM|CompTIA|ITIL)'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                cert = match.strip() if isinstance(match, str) else match[0].strip()
                if cert and len(cert) > 2 and cert not in certs:
                    certs.append(cert)
        
        return certs[:5]
    
    def extract_languages(self, text: str) -> List[str]:
        """Extract languages"""
        languages = []
        lang_list = ['english', 'amharic', 'spanish', 'french', 'german', 'arabic', 'chinese', 'japanese', 'portuguese', 'russian']
        
        for lang in lang_list:
            if lang in text.lower():
                proficiency = ""
                if 'fluent' in text.lower() and lang in text.lower():
                    proficiency = " (Fluent)"
                elif 'native' in text.lower() and lang in text.lower():
                    proficiency = " (Native)"
                elif 'basic' in text.lower() and lang in text.lower():
                    proficiency = " (Basic)"
                languages.append(lang.title() + proficiency)
        
        return languages[:5]
    
    def extract_job_titles(self, text: str) -> List[str]:
        """Extract job titles"""
        titles = []
        patterns = [
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Developer|Engineer|Designer|Analyst|Manager|Lead|Architect)',
            r'(Senior|Junior|Lead|Principal)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Developer|Engineer|Designer)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    title = ' '.join(match).strip()
                else:
                    title = match.strip()
                if title and len(title) > 3 and title not in titles:
                    titles.append(title)
        
        return list(set(titles))[:5]
    
    def parse_resume(self, file) -> Dict:
        """Complete resume parsing"""
        filename = file.filename.lower()
        
        # Extract text
        if filename.endswith('.docx'):
            text = self.extract_text_from_docx(file.file)
        elif filename.endswith('.pdf'):
            text = self.extract_text_from_pdf(file.file)
        elif filename.endswith('.txt'):
            text = self.extract_text_from_txt(file.file)
        else:
            raise Exception("Unsupported file format. Use PDF, DOCX, or TXT.")
        
        if not text or len(text.strip()) < 50:
            raise Exception("Could not extract enough text from the file.")
        
        # Extract all information
        result = {
            "filename": filename,
            "word_count": len(text.split()),
            "file_type": "document",
            # Contact info
            "email": self.extract_email(text),
            "phone": self.extract_phone(text),
            "linkedin": self.extract_linkedin(text),
            "location": self.extract_location(text),
            # Professional info
            "summary": self.extract_summary(text),
            "skills": self.extract_skills(text),
            "experience_years": self.extract_experience_years(text),
            "job_titles": self.extract_job_titles(text),
            "education": self.extract_education(text),
            "certifications": self.extract_certifications(text),
            "languages": self.extract_languages(text)
        }
        
        return result

resume_parser = ResumeParser()